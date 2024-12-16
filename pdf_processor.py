import os
import sys
import traceback
import platform
import subprocess

# PDF libraries
from pdfminer.high_level import extract_text as extract_pdf_text
import PyPDF2
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

# DOCX libraries
import docx
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Text file library
import chardet

# Google Gemini API
import google.generativeai as genai


class AdvancedDocumentTranslator:
    def __init__(self, api_key):
        """
        Initialize translator with Gemini API

        :param api_key: Google Gemini API key
        """
        genai.configure(api_key=api_key)
        self.model = genai.GenerativeModel("gemini-pro")

    def detect_encoding(self, file_path):
        """
        Detect file encoding

        :param file_path: Path to the file
        :return: Detected encoding
        """
        with open(file_path, "rb") as file:
            raw_data = file.read()
            result = chardet.detect(raw_data)
            return result["encoding"] or "utf-8"

    def extract_text(self, file_path):
        """
        Extract text from different file types

        :param file_path: Path to the file
        :return: Extracted text and additional metadata
        """
        # Validate file existence
        if not os.path.exists(file_path):
            print(f"Error: File {file_path} does not exist.")
            return None

        # Determine file type
        file_extension = os.path.splitext(file_path)[1].lower()

        try:
            # PDF extraction
            if file_extension == ".pdf":
                try:
                    # Attempt multiple PDF extraction methods
                    try:
                        text = extract_pdf_text(file_path)
                    except Exception as first_error:
                        print(f"First PDF extraction method failed: {first_error}")
                        try:
                            # Fallback method using PyPDF2
                            with open(file_path, "rb") as pdf_file:
                                pdf_reader = PyPDF2.PdfReader(pdf_file)
                                text = ""
                                for page in pdf_reader.pages:
                                    text += page.extract_text() + "\n"
                        except Exception as second_error:
                            print(
                                f"Second PDF extraction method failed: {second_error}"
                            )
                            return None

                    if not text or text.strip() == "":
                        print("No text could be extracted from the PDF.")
                        return None
                    return {"text": text, "type": "pdf"}
                except Exception as pdf_error:
                    print(f"PDF Extraction Error: {pdf_error}")
                    return None

            # DOCX extraction
            elif file_extension == ".docx":
                try:
                    doc = docx.Document(file_path)
                    paragraphs = []
                    paragraph_styles = []

                    for paragraph in doc.paragraphs:
                        if paragraph.text:
                            paragraphs.append(paragraph.text)
                            # Capture paragraph style details
                            paragraph_styles.append(
                                {
                                    "alignment": paragraph.alignment,
                                    "style": (
                                        paragraph.style.name
                                        if paragraph.style
                                        else None
                                    ),
                                    "font_name": (
                                        paragraph.runs[0].font.name
                                        if paragraph.runs
                                        else None
                                    ),
                                    "font_size": (
                                        paragraph.runs[0].font.size
                                        if paragraph.runs
                                        else None
                                    ),
                                }
                            )

                    if not paragraphs:
                        print("No text could be extracted from the DOCX.")
                        return None

                    return {
                        "text": "\n".join(paragraphs),
                        "type": "docx",
                        "styles": paragraph_styles,
                    }
                except Exception as docx_error:
                    print(f"DOCX Extraction Error: {docx_error}")
                    return None

            # TXT extraction
            elif file_extension in [".txt", ".text"]:
                try:
                    # Detect encoding
                    encoding = self.detect_encoding(file_path)

                    with open(file_path, "r", encoding=encoding) as file:
                        text = file.read()
                        if not text:
                            print("No text could be extracted from the text file.")
                            return None
                        return {"text": text, "type": "txt", "encoding": encoding}
                except Exception as txt_error:
                    print(f"Text File Extraction Error: {txt_error}")
                    return None

            else:
                print(f"Unsupported file type: {file_extension}")
                return None

        except Exception as e:
            print(f"Unexpected error processing {file_path}: {e}")
            traceback.print_exc()
            return None

    def translate_text(self, text, target_language):
        """
        Translate text using Gemini API

        :param text: Text to translate
        :param target_language: Target language
        :return: Translated text
        """
        if not text:
            print("No text to translate")
            return None

        try:
            # Handle large texts by chunking
            if len(text) > 10000:
                chunks = [text[i : i + 10000] for i in range(0, len(text), 10000)]
                translated_chunks = []

                for i, chunk in enumerate(chunks, 1):
                    print(f"Translating chunk {i}/{len(chunks)}")
                    response = self.model.generate_content(
                        f"Translate this text chunk to {target_language}:\n\n{chunk}"
                    )
                    translated_chunks.append(response.text)

                return " ".join(translated_chunks)

            else:
                prompt = f"Translate the following text to {target_language}. Preserve original formatting:\n\n{text}"
                response = self.model.generate_content(prompt)
                return response.text

        except Exception as e:
            print(f"Translation error: {e}")
            traceback.print_exc()
            return None

    def replace_document_content(self, input_file, translated_text, original_metadata):
        """
        Replace document content while preserving original formatting

        :param input_file: Original file path
        :param translated_text: Translated text
        :param original_metadata: Metadata from original document
        :return: Success status
        """
        try:
            # DOCX replacement
            if original_metadata["type"] == "docx":
                doc = docx.Document(input_file)

                # Remove all paragraphs
                while len(doc.paragraphs) > 0:
                    doc.paragraphs[0]._element.getparent().remove(
                        doc.paragraphs[0]._element
                    )

                # Add translated text with original styles
                for i, translated_paragraph in enumerate(translated_text.split("\n")):
                    # Retrieve original style
                    style = (
                        original_metadata["styles"][i]
                        if i < len(original_metadata["styles"])
                        else {}
                    )

                    # Add paragraph
                    para = doc.add_paragraph(translated_paragraph)

                    # Apply original style if available
                    if style:
                        if style["alignment"] is not None:
                            para.alignment = style["alignment"]

                        if para.runs:
                            run = para.runs[0]
                            if style["font_name"]:
                                run.font.name = style["font_name"]
                            if style["font_size"]:
                                run.font.size = style["font_size"]

                # Save the document
                output_file = input_file.replace(".docx", "_translated.docx")
                doc.save(output_file)
                print(f"Translated document saved as {output_file}")
                return True

            # TXT replacement
            elif original_metadata["type"] == "txt":
                encoding = original_metadata.get("encoding", "utf-8")
                output_file = input_file.replace(".txt", "_translated.txt")
                with open(output_file, "w", encoding=encoding) as file:
                    file.write(translated_text)
                print(f"Translated document saved as {output_file}")
                return True

            # PDF replacement (more complex)
            elif original_metadata["type"] == "pdf":
                # Create a new PDF with translated text
                output_file = input_file.replace(".pdf", "_translated.pdf")
                c = canvas.Canvas(output_file, pagesize=letter)
                width, height = letter

                # Set a default font
                pdfmetrics.registerFont(TTFont("Arial", "Arial.ttf"))
                c.setFont("Arial", 12)

                # Write translated text
                text_object = c.beginText(40, height - 40)
                for line in translated_text.split("\n"):
                    text_object.textLine(line)

                c.drawText(text_object)
                c.showPage()
                c.save()
                print(f"Translated document saved as {output_file}")
                return True

            else:
                print("Unsupported file type for replacement")
                return False

        except Exception as e:
            print(f"Error replacing document content: {e}")
            traceback.print_exc()
            return False

    def translate_document(self, input_file, target_language):
        """
        Translate document in-place

        :param input_file: Path to input file
        :param target_language: Target language
        :return: Success status
        """
        # Extract text and metadata
        document_data = self.extract_text(input_file)
        if not document_data:
            print("Failed to extract text from file")
            return False

        # Translate text
        translated_text = self.translate_text(document_data["text"], target_language)
        if not translated_text:
            print("Translation failed")
            return False

        # Replace document content
        success = self.replace_document_content(
            input_file, translated_text, document_data
        )

        if success:
            print(f"Document translated successfully: {input_file}")

        return success


def main():
    # Interactive user input
    while True:
        try:
            # Get Gemini API Key
            api_key = "AIzaSyAnEbQDilqTRSE1Bn8dqAVNhf6Ml_YyX18"
            if not api_key:
                print("API Key cannot be empty")
                continue

            # Initialize translator
            translator = AdvancedDocumentTranslator(api_key)

            # Get input file path
            input_file = input(
                "Enter the full path to the file you want to translate (PDF/DOCX/TXT): "
            ).strip()
            if not input_file:
                print("File path cannot be empty")
                continue

            # Validate file type
            supported_extensions = [".pdf", ".docx", ".txt", ".text"]
            if not any(
                input_file.lower().endswith(ext) for ext in supported_extensions
            ):
                print("Please provide a PDF, DOCX, or TXT file")
                continue

            # Get target language
            target_language = input(
                "Enter the target language (e.g., Spanish, French, German): "
            ).strip()
            if not target_language:
                print("Target language cannot be empty")
                continue

            # Translate the document
            result = translator.translate_document(input_file, target_language)

            if result:
                print("Translation completed successfully!")

            # Ask if user wants to translate another file
            another = input("Do you want to translate another file? (yes/no): ").lower()
            if another != "yes":
                break

        except Exception as e:
            print(f"An unexpected error occurred: {e}")
            traceback.print_exc()
            retry = input("Do you want to try again? (yes/no): ").lower()
            if retry != "yes":
                break


if __name__ == "__main__":
    main()
