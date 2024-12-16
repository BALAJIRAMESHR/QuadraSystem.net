import streamlit as st
import speech_recognition as sr
from translator import GeminiTranslator, LANGUAGES, LANGUAGE_CODES
import tempfile
import docx
import os
from PIL import Image
import pytesseract
from moviepy import VideoFileClip
from moviepy import TextClip, CompositeVideoClip
import io
from pdf_processor import (
    extract_pdf_text,
)  # Assuming extract_pdf_text is a function to extract text from PDFs

# Ensure session state is initialized
if "messages" not in st.session_state:
    st.session_state["messages"] = []
if "chat_name" not in st.session_state:
    st.session_state["chat_name"] = "DefaultChat"  # Default name

# Initialize translator
translator = GeminiTranslator()

# Streamlit UI Setup
st.set_page_config(page_title="Quadra Translator", page_icon="🌐")
st.title("🌐 Quadra Translator ")
st.subheader("Instant translation with language detection and file download")


# Speech Recognition Function
def recognize_speech_from_microphone():
    recognizer = sr.Recognizer()
    microphone = sr.Microphone()

    with microphone as source:
        st.info("Listening... Please speak.")
        recognizer.adjust_for_ambient_noise(source)  # Adjust to ambient noise levels
        audio = recognizer.listen(source)
        try:
            st.info("Recognizing speech...")
            text = recognizer.recognize_google(audio)  # Use Google Speech Recognition
            st.success(f"Recognized Speech: {text}")
            return text
        except sr.UnknownValueError:
            st.error("Google Speech Recognition could not understand the audio.")
            return None
        except sr.RequestError:
            st.error(
                "Could not request results from Google Speech Recognition service."
            )
            return None


# Function to extract text from various file types
def extract_text_from_file(uploaded_file):
    if uploaded_file.type == "application/pdf":
        return extract_pdf_text(uploaded_file)
    elif (
        uploaded_file.type
        == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    ):
        doc = docx.Document(uploaded_file)
        return "\n".join([para.text for para in doc.paragraphs if para.text])
    elif uploaded_file.type == "text/plain":
        return uploaded_file.getvalue().decode("utf-8")
    elif uploaded_file.type.startswith("image"):
        return extract_text_from_image(uploaded_file)
    elif uploaded_file.type.startswith("video"):
        return extract_audio_and_transcribe(uploaded_file)
    else:
        st.error("Unsupported file type")
        return None


# Function to extract text from images using pytesseract
def extract_text_from_image(image_file):
    try:
        # Open image file using PIL
        image = Image.open(image_file)
        # Use pytesseract to extract text from the image
        extracted_text = pytesseract.image_to_string(image)
        if extracted_text.strip():  # If any text was extracted
            return extracted_text
        else:
            st.error("No text found in the image.")
            return None
    except Exception as e:
        st.error(f"Error extracting text from image: {e}")
        return None


# Function to extract audio from video and transcribe it
def extract_audio_and_transcribe(video_file):
    try:
        # Save the uploaded video to a temporary file
        temp_video_path = tempfile.mktemp(suffix=".mp4")
        with open(temp_video_path, "wb") as temp_video:
            temp_video.write(video_file.getvalue())

        # Ensure video clip is opened and processed safely
        with VideoFileClip(temp_video_path) as video_clip:
            # Extract audio from video
            audio_path = tempfile.mktemp(suffix=".wav")
            video_clip.audio.write_audiofile(audio_path)

        # Use Speech Recognition to transcribe the audio
        recognizer = sr.Recognizer()
        with sr.AudioFile(audio_path) as source:
            audio_data = recognizer.record(source)
            text = recognizer.recognize_google(audio_data)

        # Clean up temporary files
        os.remove(temp_video_path)
        os.remove(audio_path)

        return text
    except Exception as e:
        print(f"Error extracting or transcribing audio: {e}")
        return None


# Function to create translated file for download
def create_translated_file(original_file, translated_text, original_type):
    if original_type == "application/pdf":
        # Create PDF with translated text
        packet = io.BytesIO()
        c = canvas.Canvas(packet, pagesize=letter)
        width, height = letter

        # Set a default font
        pdfmetrics.registerFont(TTFont("Arial", "Arial.ttf"))
        c.setFont("Arial", 12)

        # Write translated text
        text_object = c.beginText(40, height - 40)
        for line in translated_text.split("\n"):
            text_object.textLine(line)

        c.drawText(text_object)
        c.showPage()
        c.save()

        # Move to the beginning of the StringIO buffer
        packet.seek(0)
        return packet

    elif (
        original_type
        == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    ):
        # Create DOCX with translated text
        doc = docx.Document()
        for line in translated_text.split("\n"):
            doc.add_paragraph(line)

        # Save to BytesIO
        docx_file = io.BytesIO()
        doc.save(docx_file)
        docx_file.seek(0)
        return docx_file

    elif original_type == "text/plain":
        # Create text file
        return io.BytesIO(translated_text.encode("utf-8"))

    else:
        st.error("Unsupported file type for download")
        return None


# Main translation interface (Upper chat messages, bottom input)
col1, col2 = st.columns(2)

with col1:
    # Display current chat name at the top of the page
    st.markdown(f"### Current Chat: {st.session_state['chat_name']}")

    # Source language detection and display
    source_text = st.text_area(
        "Enter text to translate:",
        height=120,
        key="source_input",
        label_visibility="collapsed",
    )

    # Speech-to-Text button
    if st.button("Start Speech-to-Text Translation"):
        recognized_text = recognize_speech_from_microphone()
        if recognized_text:
            source_text = recognized_text

    # File upload with improved functionality
    uploaded_file = st.file_uploader(
        "Upload File",
        type=["txt", "pdf", "docx", "mp4", "mov", "jpg", "jpeg", "png"],
        label_visibility="collapsed",
    )

    # Process uploaded file
    translated_file = None
    if uploaded_file is not None:
        try:
            # Extract text from the uploaded file
            uploaded_text = extract_text_from_file(uploaded_file)
            if uploaded_text:
                source_text = uploaded_text
                st.success("File uploaded successfully!")
        except Exception as e:
            st.error(f"Error uploading file: {e}")

    # Detect source language
    if source_text:
        detected_lang = translator.detect_language(source_text)
        detected_language_name = LANGUAGE_CODES.get(detected_lang, "Unknown")
        st.write(f"Detected Language: {detected_language_name}")

with col2:
    # Target language selection
    target_language = st.selectbox(
        "Select target language:", list(LANGUAGES.keys()), index=1, key="target_lang"
    )

    # Perform live translation
    translated_text = None
    if source_text:
        detected_source_lang = translator.detect_language(source_text)
        target_lang_code = LANGUAGES[target_language]

        if detected_source_lang != target_lang_code:
            translated_text = translator.translate(source_text, target_lang_code)

            if translated_text:
                st.text_area(
                    f"Translated Text ({target_language}):",
                    value=translated_text,
                    height=120,
                    disabled=True,
                    label_visibility="collapsed",
                )

                # Add to history
                translator.add_to_history(
                    st.session_state["chat_name"],
                    source_text,
                    detected_source_lang,
                    target_lang_code,
                    translated_text,
                )

                # Update session state messages
                st.session_state["messages"].append(
                    {
                        "original": source_text,
                        "translation": translated_text,
                        "source_language": LANGUAGE_CODES.get(
                            detected_source_lang, "Unknown"
                        ),
                        "target_language": target_language,
                    }
                )

                # Create downloadable translated file if original file was uploaded
                if uploaded_file:
                    translated_file = create_translated_file(
                        uploaded_file, translated_text, uploaded_file.type
                    )

                    # Download button for translated file
                    if translated_file:
                        st.download_button(
                            label=f"Download Translated {uploaded_file.name}",
                            data=translated_file,
                            file_name=f"translated_{uploaded_file.name}",
                            mime=uploaded_file.type,
                        )
        else:
            st.text_area(
                f"Original Text ({detected_language_name}):",
                value=source_text,
                height=120,
                disabled=True,
                label_visibility="collapsed",
            )

# Display chat history (WhatsApp-like chat bubbles)
st.subheader("Chat-like Translation History")
if st.session_state["messages"]:
    for msg in st.session_state["messages"]:
        st.markdown(
            f"<div style='display: flex; justify-content: flex-start;'><div style='background-color: #f1f1f1; border-radius: 10px; padding: 10px; max-width: 60%; font-size: 14px;'>{msg['original']} ({msg['source_language']})</div></div>",
            unsafe_allow_html=True,
        )
        st.markdown(
            f"<div style='display: flex; justify-content: flex-end;'><div style='background-color: #4CAF50; color: white; border-radius: 10px; padding: 10px; max-width: 60%; font-size: 14px;'>{msg['translation']} ({msg['target_language']})</div></div>",
            unsafe_allow_html=True,
        )
